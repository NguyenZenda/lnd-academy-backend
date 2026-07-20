import os
import json
import re
import time
import uuid
from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Header, Depends
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, EmailStr
from huggingface_hub import InferenceClient
from typing import Optional, List
from supabase import create_client, Client

app = FastAPI()
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

HF_TOKEN = os.environ.get("HF_TOKEN", "")

# ---------------------------------------------------------------------------
# Supabase setup (Auth + Postgres + Storage)
# ---------------------------------------------------------------------------
SUPABASE_URL = os.environ.get("SUPABASE_URL", "")
SUPABASE_SERVICE_KEY = os.environ.get("SUPABASE_SERVICE_KEY", "")
TEACHER_INVITE_CODE = os.environ.get("TEACHER_INVITE_CODE", "")
DOCUMENTS_BUCKET = "documents"

supabase: Optional[Client] = (
    create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)
    if SUPABASE_URL and SUPABASE_SERVICE_KEY
    else None
)

ALLOWED_CATEGORIES = {"ielts", "hsca", "thpt"}
ALLOWED_ROLES = {"student", "teacher"}


def require_supabase():
    if not supabase:
        raise HTTPException(status_code=500, detail="Auth/Database chưa được cấu hình (thiếu SUPABASE_URL/SUPABASE_SERVICE_KEY)")


def get_current_user(authorization: Optional[str] = Header(None)):
    require_supabase()
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Thiếu token đăng nhập")
    token = authorization.split(" ", 1)[1]
    try:
        user_resp = supabase.auth.get_user(token)
        user = user_resp.user
        if not user:
            raise HTTPException(status_code=401, detail="Token không hợp lệ")
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(status_code=401, detail="Token không hợp lệ hoặc đã hết hạn")

    profile = supabase.table("profiles").select("*").eq("id", user.id).single().execute()
    if not profile.data:
        raise HTTPException(status_code=401, detail="Không tìm thấy hồ sơ người dùng")
    return {"id": user.id, "email": user.email, **profile.data}


def get_optional_user(authorization: Optional[str] = Header(None)):
    """Giong get_current_user nhung khong bao gio raise loi - tra None neu chua dang nhap/token sai.
    Dung cho cac endpoint muon "mo cho khach" nhung van nhan biet duoc user da dang nhap (VD nop bai thi)."""
    if not supabase or not authorization or not authorization.startswith("Bearer "):
        return None
    token = authorization.split(" ", 1)[1]
    try:
        user_resp = supabase.auth.get_user(token)
        user = user_resp.user
        if not user:
            return None
        profile = supabase.table("profiles").select("*").eq("id", user.id).single().execute()
        if not profile.data:
            return None
        return {"id": user.id, "email": user.email, **profile.data}
    except Exception:
        return None


# ---------------------------------------------------------------------------
# Auth
# ---------------------------------------------------------------------------
class RegisterRequest(BaseModel):
    email: EmailStr
    password: str
    full_name: str
    role: str  # "student" | "teacher"
    invite_code: Optional[str] = None


class LoginRequest(BaseModel):
    email: EmailStr
    password: str


@app.post("/auth/register")
def register(req: RegisterRequest):
    require_supabase()
    if req.role not in ALLOWED_ROLES:
        raise HTTPException(status_code=400, detail="Vai trò không hợp lệ")
    if len(req.password) < 6:
        raise HTTPException(status_code=400, detail="Mật khẩu phải có ít nhất 6 ký tự")

    if req.role == "teacher":
        if not TEACHER_INVITE_CODE:
            raise HTTPException(status_code=500, detail="Hệ thống chưa cấu hình mã mời giáo viên")
        if req.invite_code != TEACHER_INVITE_CODE:
            raise HTTPException(status_code=403, detail="Mã mời giáo viên không đúng")

    try:
        auth_resp = supabase.auth.admin.create_user({
            "email": req.email,
            "password": req.password,
            "email_confirm": True,
        })
        user_id = auth_resp.user.id
        supabase.table("profiles").insert({
            "id": user_id,
            "full_name": req.full_name,
            "role": req.role,
        }).execute()
        return {"status": "ok", "message": "Đăng ký thành công, mời bạn đăng nhập"}
    except Exception as e:
        msg = str(e)
        if "already" in msg.lower() or "duplicate" in msg.lower():
            msg = "Email này đã được đăng ký"
        raise HTTPException(status_code=400, detail=msg)


@app.post("/auth/login")
def login(req: LoginRequest):
    require_supabase()
    try:
        resp = supabase.auth.sign_in_with_password({"email": req.email, "password": req.password})
        profile = supabase.table("profiles").select("*").eq("id", resp.user.id).single().execute()
        return {
            "access_token": resp.session.access_token,
            "refresh_token": resp.session.refresh_token,
            "expires_at": resp.session.expires_at,
            "user": {"id": resp.user.id, "email": resp.user.email, **(profile.data or {})},
        }
    except Exception:
        raise HTTPException(status_code=401, detail="Email hoặc mật khẩu không đúng")


class RefreshRequest(BaseModel):
    refresh_token: str


@app.post("/auth/refresh")
def refresh_session(req: RefreshRequest):
    require_supabase()
    try:
        resp = supabase.auth.refresh_session(req.refresh_token)
        return {
            "access_token": resp.session.access_token,
            "refresh_token": resp.session.refresh_token,
            "expires_at": resp.session.expires_at,
        }
    except Exception:
        raise HTTPException(status_code=401, detail="Không thể làm mới phiên đăng nhập, vui lòng đăng nhập lại")


@app.get("/me")
def me(user=Depends(get_current_user)):
    return user


# ---------------------------------------------------------------------------
# Documents (IELTS / HSCA / THPT libraries)
# ---------------------------------------------------------------------------
@app.get("/documents")
def list_documents(category: str, skill: Optional[str] = None, levels: Optional[str] = None):
    require_supabase()
    if category not in ALLOWED_CATEGORIES:
        raise HTTPException(status_code=400, detail="Danh mục không hợp lệ")
    q = supabase.table("documents").select("*").eq("category", category).order("created_at", desc=True)
    if skill:
        q = q.eq("skill", skill)
    docs = q.execute().data
    if levels:
        wanted = set(levels.split(","))
        docs = [d for d in docs if wanted & set(d.get("cefr_levels") or [])]
    return docs


@app.post("/documents/upload")
async def upload_document(
    category: str = Form(...),
    skill: str = Form(""),
    cefr_levels: str = Form(""),  # chuoi JSON list, vd ["B1","B2"]
    title: str = Form(...),
    file: UploadFile = File(...),
    user=Depends(get_current_user),
):
    if user["role"] != "teacher":
        raise HTTPException(status_code=403, detail="Chỉ giáo viên mới được đăng tài liệu")
    if category not in ALLOWED_CATEGORIES:
        raise HTTPException(status_code=400, detail="Danh mục không hợp lệ")

    try:
        levels_list = json.loads(cefr_levels) if cefr_levels else []
    except json.JSONDecodeError:
        levels_list = []

    file_bytes = await file.read()
    if len(file_bytes) > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File vượt quá 20MB")

    ext = os.path.splitext(file.filename or "")[1]
    storage_path = f"{category}/{uuid.uuid4().hex}{ext}"

    try:
        supabase.storage.from_(DOCUMENTS_BUCKET).upload(
            storage_path,
            file_bytes,
            {"content-type": file.content_type or "application/octet-stream"},
        )
        public_url = supabase.storage.from_(DOCUMENTS_BUCKET).get_public_url(storage_path)

        row = {
            "title": title,
            "category": category,
            "skill": skill,
            "cefr_levels": levels_list,
            "file_url": public_url,
            "file_name": file.filename,
            "uploaded_by": user["id"],
            "uploader_name": user.get("full_name", ""),
        }
        supabase.table("documents").insert(row).execute()
        return {"status": "ok", "file_url": public_url}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/documents/{doc_id}")
def delete_document(doc_id: str, user=Depends(get_current_user)):
    if user["role"] != "teacher":
        raise HTTPException(status_code=403, detail="Chỉ giáo viên mới được xoá tài liệu")
    doc = supabase.table("documents").select("*").eq("id", doc_id).single().execute()
    if not doc.data:
        raise HTTPException(status_code=404, detail="Không tìm thấy tài liệu")
    if doc.data["uploaded_by"] != user["id"]:
        raise HTTPException(status_code=403, detail="Bạn chỉ có thể xoá tài liệu do chính mình đăng")
    supabase.table("documents").delete().eq("id", doc_id).execute()
    return {"status": "ok"}


# ---------------------------------------------------------------------------
# IELTS Writing grader (unchanged from existing production logic)
# ---------------------------------------------------------------------------
TASK1_DESCRIPTORS = """
IELTS WRITING TASK 1 - OFFICIAL BAND DESCRIPTORS (Updated May 2023)
Each criterion is scored as a WHOLE NUMBER only (1-9). No decimals. No .5.

TASK ACHIEVEMENT (TA):
9: All requirements fully and appropriately satisfied. Extremely rare lapses.
8: Covers all requirements appropriately, relevantly, sufficiently. Key features skilfully selected, presented, highlighted, illustrated. Occasional omissions.
7: Covers requirements. Content relevant and accurate with few omissions. Key features covered and clearly highlighted. Clear overview, data appropriately categorised, main trends identified.
6: Focuses on requirements. Key features covered and adequately highlighted. Overview attempted. Some irrelevant/inaccurate info may occur. Some details missing or excessive.
5: Generally addresses requirements. Key features not adequately covered. Recounting mainly mechanical. May be no data to support description. Tendency to focus on details without bigger picture.
4: Attempts to address task. Few key features selected. Key features may be irrelevant, repetitive, inaccurate. Format may be inappropriate.
3: Does not address requirements (possibly misunderstood data/diagram). Key features largely irrelevant. Limited information, used repetitively.
2: Content barely relates to task.
1: Content wholly unrelated to task (or 20 words or fewer).

COHERENCE & COHESION (CC):
9: Message followed effortlessly. Cohesion rarely attracts attention. Minimal lapses. Paragraphing skilfully managed.
8: Message followed with ease. Ideas logically sequenced, cohesion well managed. Occasional lapses. Paragraphing used sufficiently and appropriately.
7: Ideas logically organised, clear progression. Few lapses. Cohesive devices used flexibly but with some inaccuracies or over/under use.
6: Generally arranged coherently, clear overall progression. Cohesive devices used to some good effect but cohesion may be faulty or mechanical. Reference/substitution may lack flexibility.
5: Organisation evident but not wholly logical, may lack overall progression. Relationship of ideas can be followed but sentences not fluently linked. Limited/overuse of cohesive devices.
4: Ideas evident but not arranged coherently, no clear progression. Relationships unclear/inadequately marked. Inaccurate use or lack of substitution/referencing.
3: No apparent logical organisation. Minimal use of sequencers/cohesive devices. Difficulty identifying referencing.
2: Little relevant message. Little evidence of control of organisational features.
1: Writing fails to communicate any message.

LEXICAL RESOURCE (LR):
9: Full flexibility and precise use evident. Wide range used accurately and appropriately. Very natural and sophisticated control. Minor errors extremely rare.
8: Wide resource fluently and flexibly used to convey precise meanings. Skilful use of uncommon/idiomatic items. Occasional errors in spelling/word formation with minimal impact.
7: Sufficient for some flexibility and precision. Some ability to use less common/idiomatic items. Awareness of style and collocation, though inappropriacies occur. Few errors in spelling/word formation.
6: Generally adequate and appropriate. Meaning generally clear despite restricted range or lack of precision. Some errors in spelling/word formation but do not impede communication.
5: Limited but minimally adequate. Simple vocabulary used accurately but range doesn't permit much variation. Frequent lapses in appropriacy. Errors in spelling/word formation may cause difficulty.
4: Limited and inadequate for task. Basic vocabulary, may be used repetitively. Inappropriate word choice/errors in word formation may impede meaning.
3: Inadequate. Possible over-dependence on memorised language. Control of word choice/spelling very limited, errors predominate and may severely impede meaning.
2: Extremely limited. Few recognisable strings apart from memorised phrases. No apparent control of word formation/spelling.
1: No resource apparent except a few isolated words.

GRAMMATICAL RANGE & ACCURACY (GRA):
9: Wide range of structures used with full flexibility and control. Punctuation and grammar used appropriately throughout. Minor errors extremely rare.
8: Wide range of structures flexibly and accurately used. Majority of sentences error-free, punctuation well managed. Occasional non-systematic errors with minimal impact.
7: Variety of complex structures used with some flexibility and accuracy. Grammar and punctuation generally well controlled. Error-free sentences frequent. Few errors persist but don't impede communication.
6: Mix of simple and complex sentence forms but flexibility limited. More complex structures not as accurate as simple. Errors in grammar and punctuation occur but rarely impede communication.
5: Range of structures limited and rather repetitive. Complex sentences attempted but tend to be faulty. Grammatical errors may be frequent and cause difficulty. Punctuation may be faulty.
4: Very limited range of structures. Subordinate clauses rare, simple sentences predominate. Grammatical errors frequent and may impede meaning. Punctuation often faulty.
3: Sentence forms attempted but errors in grammar and punctuation predominate. Prevents most meaning from coming through. Length may be insufficient.
2: Little or no evidence of sentence forms except in memorised phrases.
1: No rateable language evident.

OVERALL BAND CALCULATION for Task 1:
Average = (TA + CC + LR + GRA) / 4
Round to nearest 0.5: if average ends in .25 round up to .5, if ends in .75 round up to next whole number.
Example: (6+6+7+6)/4 = 6.25 → Overall Band 6.5
"""

TASK2_DESCRIPTORS = """
IELTS WRITING TASK 2 - OFFICIAL BAND DESCRIPTORS (Updated May 2023)
Each criterion is scored as a WHOLE NUMBER only (1-9). No decimals. No .5.

TASK RESPONSE (TR):
9: Prompt appropriately addressed and explored in depth. Clear and fully developed position which directly answers question. Ideas relevant, fully extended and well supported.
8: Prompt appropriately and sufficiently addressed. Clear and well-developed position. Ideas relevant, well extended and supported. Occasional omissions.
7: Main parts of prompt appropriately addressed. Clear and developed position. Main ideas extended and supported but may over-generalise or lack focus/precision.
6: Main parts addressed (some more fully than others). Position directly relevant to prompt but conclusions may be unclear/unjustified. Main ideas relevant but some insufficiently developed.
5: Main parts incompletely addressed. Format may be inappropriate. Position expressed but development not always clear. Some main ideas put forward but limited and not sufficiently developed.
4: Prompt tackled minimally or tangentially. Position discernible but hard to find. Main ideas difficult to identify. Large parts may be repetitive.
3: No part adequately addressed or prompt misunderstood. No relevant position. Few ideas, may be irrelevant or insufficiently developed.
2: Content barely related to prompt. No position identifiable.
1: Content wholly unrelated to prompt (or 20 words or fewer).

COHERENCE & COHESION (CC):
9: Message followed effortlessly. Cohesion rarely attracts attention. Minimal lapses. Paragraphing skilfully managed.
8: Message followed with ease. Ideas logically sequenced, cohesion well managed. Occasional lapses. Paragraphing used sufficiently and appropriately.
7: Ideas logically organised, clear progression throughout. Few minor lapses. Cohesive devices used flexibly but with some inaccuracies or over/under use. Paragraphing generally effective.
6: Generally arranged coherently, clear overall progression. Cohesive devices used to some good effect but cohesion may be faulty or mechanical. Paragraphing may not always be logical.
5: Organisation evident but not wholly logical, may lack overall progression. Relationship of ideas can be followed but sentences not fluently linked. Paragraphing may be inadequate or missing.
4: Ideas evident but not arranged coherently, no clear progression. Relationships unclear. May be no paragraphing or no clear main topic within paragraphs.
3: No apparent logical organisation. Minimal use of sequencers/cohesive devices. Difficulty identifying referencing. Any paragraphing attempts are unhelpful.
2: Little relevant message. Little evidence of control of organisational features.
1: Writing fails to communicate any message.

LEXICAL RESOURCE (LR):
9: Full flexibility and precise use widely evident. Wide range used accurately and appropriately. Very natural and sophisticated control. Minor errors extremely rare.
8: Wide resource fluently and flexibly used to convey precise meanings. Skilful use of uncommon/idiomatic items. Occasional errors with minimal impact.
7: Sufficient for some flexibility and precision. Some ability to use less common/idiomatic items. Awareness of style and collocation, though inappropriacies occur. Few errors.
6: Generally adequate and appropriate. Meaning generally clear despite restricted range or lack of precision. Some errors but do not impede communication.
5: Limited but minimally adequate. Simple vocabulary used accurately but range doesn't permit much variation. Frequent lapses in appropriacy. Errors may cause difficulty.
4: Limited and inadequate for task. Basic vocabulary, may be used repetitively. Errors in word formation/spelling may impede meaning.
3: Inadequate. Possible over-dependence on memorised language. Control very limited, errors predominate and may severely impede meaning.
2: Extremely limited. Few recognisable strings. No apparent control.
1: No resource apparent.

GRAMMATICAL RANGE & ACCURACY (GRA):
9: Wide range of structures used with full flexibility and control. Grammar and punctuation used appropriately throughout. Minor errors extremely rare.
8: Wide range flexibly and accurately used. Majority of sentences error-free. Occasional non-systematic errors with minimal impact.
7: Variety of complex structures with some flexibility and accuracy. Grammar and punctuation generally well controlled. Error-free sentences frequent. Few errors persist.
6: Mix of simple and complex forms but flexibility limited. More complex structures not as accurate. Errors in grammar and punctuation occur but rarely impede communication.
5: Range limited and rather repetitive. Complex sentences tend to be faulty. Grammatical errors may be frequent and cause difficulty. Punctuation may be faulty.
4: Very limited range. Simple sentences predominate. Grammatical errors frequent and may impede meaning. Punctuation often faulty.
3: Sentence forms attempted but errors predominate. Prevents most meaning from coming through.
2: Little or no evidence of sentence forms.
1: No rateable language evident.

OVERALL BAND CALCULATION for Task 2:
Average = (TR + CC + LR + GRA) / 4
Round to nearest 0.5: if average ends in .25 round up to .5, if ends in .75 round up to next whole number.
Example: (6+7+6+6)/4 = 6.25 → Overall Band 6.5
Example: (7+7+7+6)/4 = 6.75 → Overall Band 7.0
"""


class GradeRequest(BaseModel):
    essay: str
    prompt: str
    task_type: str
    image_base64: Optional[str] = None


@app.get("/")
def root():
    return {"status": "LND Academy API is running"}


@app.post("/grade")
def grade_essay(req: GradeRequest):
    if not HF_TOKEN:
        raise HTTPException(status_code=500, detail="HF_TOKEN not configured")
    if len(req.essay.strip()) < 50:
        raise HTTPException(status_code=400, detail="Essay is too short")

    client = InferenceClient(provider="nscale", api_key=HF_TOKEN)

    if req.task_type == "Task 1":
        criterion1_name = "Task Achievement (TA)"
        descriptors = TASK1_DESCRIPTORS
    else:
        criterion1_name = "Task Response (TR)"
        descriptors = TASK2_DESCRIPTORS

    system_prompt = (
        "You are a certified IELTS examiner with 10+ years of experience. "
        "You must grade strictly and accurately following the official IELTS Band Descriptors below.\n\n"
        + descriptors +
        "\n\nCRITICAL RULES:\n"
        "1. Each criterion score MUST be a WHOLE NUMBER (1-9). NEVER use decimals like 6.5 for individual criteria.\n"
        "2. Only the overall_band can be x.0 or x.5 (calculated by averaging the 4 criteria and rounding to nearest 0.5).\n"
        "3. Read the band descriptors carefully for EACH criterion before scoring.\n"
        "4. Be strict and accurate — do not inflate scores.\n"
        "5. Respond ONLY in this exact JSON format, no extra text, no markdown:\n"
        "{\n"
        f'  "criterion1_name": "{criterion1_name}",\n'
        '  "criterion1_score": 6,\n'
        '  "criterion1_feedback": "Specific feedback referencing the band descriptors...",\n'
        '  "criterion2_name": "Coherence & Cohesion",\n'
        '  "criterion2_score": 6,\n'
        '  "criterion2_feedback": "Specific feedback referencing the band descriptors...",\n'
        '  "criterion3_name": "Lexical Resource",\n'
        '  "criterion3_score": 7,\n'
        '  "criterion3_feedback": "Specific feedback referencing the band descriptors...",\n'
        '  "criterion4_name": "Grammatical Range & Accuracy",\n'
        '  "criterion4_score": 6,\n'
        '  "criterion4_feedback": "Specific feedback referencing the band descriptors...",\n'
        '  "overall_band": 6.5,\n'
        '  "overall_feedback": "Comprehensive examiner feedback on the essay overall..."\n'
        "}"
    )

    user_content = f"EXAM PROMPT:\n{req.prompt}\n\nSTUDENT ESSAY:\n{req.essay}"
    user_message = f"Grade this IELTS Writing {req.task_type} using the official band descriptors provided:\n\n{user_content}\n/no_think"

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_message}
    ]

    try:
        response = client.chat.completions.create(
            model="Qwen/Qwen3-32B",
            messages=messages,
            max_tokens=2000,
        )
        raw = response.choices[0].message.content.strip()
        raw = re.sub(r'<think>.*?</think>', '', raw, flags=re.DOTALL).strip()
        if "```" in raw:
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        match = re.search(r'\{.*\}', raw, re.DOTALL)
        if match:
            raw = match.group(0)
        result = json.loads(raw.strip())

        for key in ["criterion1_score", "criterion2_score", "criterion3_score", "criterion4_score"]:
            if key in result:
                result[key] = int(round(result[key]))

        scores = [result["criterion1_score"], result["criterion2_score"],
                  result["criterion3_score"], result["criterion4_score"]]
        avg = sum(scores) / 4
        result["overall_band"] = round(avg * 2) / 2

        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ---------------------------------------------------------------------------
# Live chat assistant (site-wide widget)
# ---------------------------------------------------------------------------
CHAT_SYSTEM_PROMPT = """Bạn là trợ lý ảo của LND Academy — trung tâm tiếng Anh học thuật chất lượng cao tại tỉnh Lâm Đồng.

VỀ TRUNG TÂM:
- Sứ mệnh: đồng hành và nâng tầm năng lực Anh ngữ học thuật cho học sinh Lâm Đồng, biến tiếng Anh thành môn học được yêu thích chứ không phải nỗi sợ, giúp học sinh thật sự dùng được ngôn ngữ để giao tiếp chứ không chỉ học để lấy điểm.
- Triết lý đào tạo: "Trust the process. See your progress." — kiên trì theo lộ trình khoa học, cá nhân hóa.

LỘ TRÌNH ĐÀO TẠO IELTS:
- Foundation: đầu vào 3.0 → đầu ra 5.5. Xây nền ngữ pháp, phát âm cốt lõi.
- Intermediate: đầu vào 5.5 → đầu ra 6.5. Phát triển tư duy nghị luận, triển khai đều 4 kỹ năng.
- Advanced: đầu vào 6.5 → đầu ra 7.0+. Tối ưu chiến thuật làm bài, hướng điểm số xuất sắc.
- Đặc điểm chung: sĩ số nhỏ cực hạn, tương tác sâu sát, sửa bài 1:1 chuyên sâu.

CHƯƠNG TRÌNH H-SCA (ôn thi Đánh giá năng lực chuyên biệt của Đại học Sư phạm TP.HCM):
- H-SCA Prep Course: lộ trình dài hơi, hệ thống hóa kiến thức trọng tâm ngữ pháp & từ vựng.
- H-SCA Crash Course: tổng ôn và luyện đề cấp tốc trong thời gian ngắn trước kỳ thi.

TÍNH NĂNG TRÊN WEBSITE:
- "Chấm bài Writing": chấm điểm IELTS Writing Task 1/Task 2 tự động bằng AI theo Band Descriptors chính thức, có nhận xét chi tiết 4 tiêu chí — dùng miễn phí, không cần đăng nhập.
- Thư viện tài liệu IELTS / HSCA / THPT: học sinh xem và tải miễn phí; giáo viên (đã xác minh qua mã mời) mới đăng tài liệu mới được.
- Cần đăng ký tài khoản để giáo viên đăng tài liệu; học sinh có thể xem tài liệu mà không cần đăng nhập.

ĐỊNH HƯỚNG SẮP RA MẮT: Ôn luyện Chuyên Anh (thi vào THPT chuyên), Bồi dưỡng Học sinh Giỏi cấp Tỉnh/Khu vực.

LIÊN HỆ:
- Hotline/Zalo: 0389 339 171 (Võ Sỹ Đồng) hoặc 0783 630 468 (Trần Long Nguyên)
- Email: lndacademy.work@gmail.com

QUY TẮC TRẢ LỜI:
- Trả lời tiếng Việt, thân thiện, ngắn gọn (2-5 câu), đi thẳng trọng tâm câu hỏi.
- Với câu hỏi học thuật (ngữ pháp, từ vựng, cấu trúc bài thi IELTS...), giải thích rõ ràng, dễ hiểu, có ví dụ ngắn nếu cần.
- Với câu hỏi về tuyển sinh (học phí cụ thể, lịch khai giảng, ưu đãi...) mà không có trong dữ liệu trên, KHÔNG bịa số liệu — hướng dẫn liên hệ hotline/Zalo/email để được tư vấn chính xác.
- Không tự ý cam kết học phí, ưu đãi, cam kết đầu ra cụ thể nếu không có trong dữ liệu trên."""


class ChatMessage(BaseModel):
    role: str
    content: str


class ChatRequest(BaseModel):
    messages: list[ChatMessage]


@app.post("/chat")
def chat(req: ChatRequest):
    if not HF_TOKEN:
        raise HTTPException(status_code=500, detail="HF_TOKEN not configured")
    if not req.messages:
        raise HTTPException(status_code=400, detail="Thiếu tin nhắn")

    client = InferenceClient(provider="nscale", api_key=HF_TOKEN)

    history = [{"role": m.role, "content": m.content} for m in req.messages[-12:]]
    messages = [{"role": "system", "content": CHAT_SYSTEM_PROMPT}] + history
    if messages[-1]["role"] == "user":
        messages[-1]["content"] += "\n/no_think"

    try:
        response = client.chat.completions.create(
            model="Qwen/Qwen3-32B",
            messages=messages,
            max_tokens=600,
        )
        raw = response.choices[0].message.content.strip()
        raw = re.sub(r'<think>.*?</think>', '', raw, flags=re.DOTALL).strip()
        return {"reply": raw}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))



# ---------------------------------------------------------------------------
# Exam bank (Reading/Listening, MCQ + short answer) — admin uploads, students take
# ---------------------------------------------------------------------------
import io
from docx import Document as DocxDocument
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload


def require_admin(user):
    if user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Chỉ admin mới được thực hiện thao tác này")


class ExamQuestionIn(BaseModel):
    question_number: int
    question_type: str  # "mcq" | "short_answer"
    question_text: str
    options: Optional[List[str]] = None
    correct_answer: str
    explanation: Optional[str] = None
    section_label: Optional[str] = None
    group_ref: Optional[str] = None  # khop voi ExamGroupIn.temp_id, null neu cau doc lap
    part_ref: Optional[str] = None  # khop voi ExamPartIn.temp_id, null neu de khong chia Part


@app.post("/exams/parse")
async def parse_exam(
    skill: str = Form(...),
    file: UploadFile = File(...),
    user=Depends(get_current_user),
):
    """Trich xuat text tho tu file docx, KHONG dung AI, KHONG tu doan cau hoi/dap an.
    Admin tu nhap tung cau hoi + dap an thu cong o buoc sau."""
    require_admin(user)
    if skill not in ("reading", "listening"):
        raise HTTPException(status_code=400, detail="Kỹ năng không hợp lệ")
    if not (file.filename or "").lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Chỉ hỗ trợ file .docx")

    content = await file.read()
    try:
        doc = DocxDocument(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        raw_text = "\n".join(paragraphs)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Không đọc được file docx: {e}")

    if len(raw_text.strip()) < 10:
        raise HTTPException(status_code=400, detail="File docx không có nội dung hoặc quá ngắn")

    return {"passage_text": raw_text, "questions": []}


class ExamGroupIn(BaseModel):
    temp_id: str  # id tam client tu sinh, dung de gan question_ref -> group that sau khi tao
    label: str
    instructions_html: str = ""
    word_bank: Optional[List[str]] = None
    order_index: int = 0
    part_ref: Optional[str] = None  # khop voi ExamPartIn.temp_id, null neu de khong chia Part


class ExamPartIn(BaseModel):
    temp_id: str
    part_number: int
    title: Optional[str] = None
    passage_text: str = ""
    audio_url: Optional[str] = None
    order_index: int = 0


class ExamCreateRequest(BaseModel):
    title: str
    skill: str
    passage_text: str
    audio_url: Optional[str] = None
    questions: List[ExamQuestionIn]
    groups: Optional[List[ExamGroupIn]] = None
    parts: Optional[List[ExamPartIn]] = None
    is_full_test: bool = True


GOOGLE_SERVICE_ACCOUNT_JSON = os.environ.get("GOOGLE_SERVICE_ACCOUNT_JSON", "")
AUDIO_DRIVE_FOLDER_ID = os.environ.get("AUDIO_DRIVE_FOLDER_ID", "")


def _get_drive_service():
    if not GOOGLE_SERVICE_ACCOUNT_JSON:
        raise HTTPException(status_code=500, detail="Chưa cấu hình GOOGLE_SERVICE_ACCOUNT_JSON")
    try:
        info = json.loads(GOOGLE_SERVICE_ACCOUNT_JSON)
    except json.JSONDecodeError:
        raise HTTPException(status_code=500, detail="GOOGLE_SERVICE_ACCOUNT_JSON không phải JSON hợp lệ")
    creds = service_account.Credentials.from_service_account_info(
        info, scopes=["https://www.googleapis.com/auth/drive.file"]
    )
    return build("drive", "v3", credentials=creds)


@app.post("/exams/upload-audio")
async def upload_exam_audio(file: UploadFile = File(...), user=Depends(get_current_user)):
    require_admin(user)
    if not (file.filename or "").lower().endswith((".mp3", ".wav", ".m4a", ".ogg")):
        raise HTTPException(status_code=400, detail="Chỉ hỗ trợ file âm thanh (.mp3, .wav, .m4a, .ogg)")
    if not AUDIO_DRIVE_FOLDER_ID:
        raise HTTPException(status_code=500, detail="Chưa cấu hình AUDIO_DRIVE_FOLDER_ID")

    file_bytes = await file.read()
    if len(file_bytes) > 100 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File âm thanh vượt quá 100MB")

    try:
        service = _get_drive_service()
        media = MediaIoBaseUpload(
            io.BytesIO(file_bytes),
            mimetype=file.content_type or "audio/mpeg",
            resumable=False,
        )
        file_metadata = {
            "name": f"{uuid.uuid4().hex}_{file.filename}",
            "parents": [AUDIO_DRIVE_FOLDER_ID],
        }
        created = service.files().create(body=file_metadata, media_body=media, fields="id").execute()
        file_id = created["id"]

        # Cho phep bat ky ai co link deu nghe/tai duoc (khong can dang nhap Google)
        service.permissions().create(
            fileId=file_id, body={"type": "anyone", "role": "reader"}
        ).execute()

        direct_url = f"https://drive.google.com/uc?export=download&id={file_id}"
        return {"audio_url": direct_url}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Upload lên Google Drive thất bại: {e}")


def _insert_parts_groups_questions(exam_id, req):
    """Dung chung cho create va update: tao parts -> groups -> questions, noi dung temp_id/group_ref/part_ref."""
    part_id_map = {}
    if req.parts:
        part_rows = [{
            "exam_id": exam_id,
            "part_number": p.part_number,
            "title": p.title,
            "passage_text": p.passage_text,
            "audio_url": p.audio_url,
            "order_index": p.order_index,
        } for p in req.parts]
        part_res = supabase.table("exam_parts").insert(part_rows).execute()
        for temp_part, real_part in zip(req.parts, part_res.data):
            part_id_map[temp_part.temp_id] = real_part["id"]

    group_id_map = {}
    if req.groups:
        group_rows = [{
            "exam_id": exam_id,
            "label": g.label,
            "instructions_html": g.instructions_html,
            "word_bank": g.word_bank,
            "order_index": g.order_index,
            "part_id": part_id_map.get(g.part_ref) if g.part_ref else None,
        } for g in req.groups]
        group_res = supabase.table("exam_groups").insert(group_rows).execute()
        for temp_group, real_group in zip(req.groups, group_res.data):
            group_id_map[temp_group.temp_id] = real_group["id"]

    question_rows = [{
        "exam_id": exam_id,
        "question_number": q.question_number,
        "question_type": q.question_type,
        "question_text": q.question_text,
        "options": q.options,
        "correct_answer": q.correct_answer,
        "explanation": q.explanation,
        "section_label": q.section_label,
        "group_id": group_id_map.get(q.group_ref) if q.group_ref else None,
        "part_id": part_id_map.get(q.part_ref) if q.part_ref else None,
    } for q in req.questions]
    supabase.table("exam_questions").insert(question_rows).execute()


@app.post("/exams")
def create_exam(req: ExamCreateRequest, user=Depends(get_current_user)):
    require_admin(user)
    if req.skill not in ("reading", "listening"):
        raise HTTPException(status_code=400, detail="Kỹ năng không hợp lệ")
    if not req.questions:
        raise HTTPException(status_code=400, detail="Đề thi cần có ít nhất 1 câu hỏi")

    exam_row = {
        "title": req.title,
        "skill": req.skill,
        "passage_text": req.passage_text,
        "audio_url": req.audio_url,
        "is_full_test": req.is_full_test,
        "created_by": user["id"],
    }
    exam_res = supabase.table("exams").insert(exam_row).execute()
    exam_id = exam_res.data[0]["id"]

    _insert_parts_groups_questions(exam_id, req)

    return {"status": "ok", "exam_id": exam_id}


@app.get("/exams")
def list_exams(skill: Optional[str] = None):
    require_supabase()
    q = supabase.table("exams").select("id,title,skill,is_full_test,created_at").order("created_at", desc=True)
    if skill:
        q = q.eq("skill", skill)
    return q.execute().data


@app.get("/exams/{exam_id}")
def get_exam(exam_id: str):
    require_supabase()
    exam = supabase.table("exams").select("*").eq("id", exam_id).single().execute()
    if not exam.data:
        raise HTTPException(status_code=404, detail="Không tìm thấy đề thi")
    questions = (
        supabase.table("exam_questions").select("*").eq("exam_id", exam_id)
        .order("question_number").execute().data
    )
    for q in questions:
        q.pop("correct_answer", None)  # an dap an, khong lo lo cho hoc sinh
    groups = (
        supabase.table("exam_groups").select("*").eq("exam_id", exam_id)
        .order("order_index").execute().data
    )
    parts = (
        supabase.table("exam_parts").select("*").eq("exam_id", exam_id)
        .order("order_index").execute().data
    )
    return {"exam": exam.data, "questions": questions, "groups": groups, "parts": parts}


@app.get("/exams/{exam_id}/edit")
def get_exam_for_edit(exam_id: str, user=Depends(get_current_user)):
    """Giong get_exam nhung KHONG an dap an - chi admin dung de sua de thi."""
    require_admin(user)
    exam = supabase.table("exams").select("*").eq("id", exam_id).single().execute()
    if not exam.data:
        raise HTTPException(status_code=404, detail="Không tìm thấy đề thi")
    questions = (
        supabase.table("exam_questions").select("*").eq("exam_id", exam_id)
        .order("question_number").execute().data
    )
    groups = (
        supabase.table("exam_groups").select("*").eq("exam_id", exam_id)
        .order("order_index").execute().data
    )
    parts = (
        supabase.table("exam_parts").select("*").eq("exam_id", exam_id)
        .order("order_index").execute().data
    )
    return {"exam": exam.data, "questions": questions, "groups": groups, "parts": parts}


@app.put("/exams/{exam_id}")
def update_exam(exam_id: str, req: ExamCreateRequest, user=Depends(get_current_user)):
    require_admin(user)
    if req.skill not in ("reading", "listening"):
        raise HTTPException(status_code=400, detail="Kỹ năng không hợp lệ")
    if not req.questions:
        raise HTTPException(status_code=400, detail="Đề thi cần có ít nhất 1 câu hỏi")

    existing = supabase.table("exams").select("id").eq("id", exam_id).single().execute()
    if not existing.data:
        raise HTTPException(status_code=404, detail="Không tìm thấy đề thi")

    supabase.table("exams").update({
        "title": req.title,
        "skill": req.skill,
        "passage_text": req.passage_text,
        "audio_url": req.audio_url,
        "is_full_test": req.is_full_test,
    }).eq("id", exam_id).execute()

    # Xoa cau hoi (con) truoc, roi nhom, roi parts (cha), sau do tao lai tu dau
    supabase.table("exam_questions").delete().eq("exam_id", exam_id).execute()
    supabase.table("exam_groups").delete().eq("exam_id", exam_id).execute()
    supabase.table("exam_parts").delete().eq("exam_id", exam_id).execute()

    _insert_parts_groups_questions(exam_id, req)

    return {"status": "ok", "exam_id": exam_id}


@app.delete("/exams/{exam_id}")
def delete_exam(exam_id: str, user=Depends(get_current_user)):
    require_admin(user)
    supabase.table("exam_questions").delete().eq("exam_id", exam_id).execute()
    supabase.table("exam_groups").delete().eq("exam_id", exam_id).execute()
    supabase.table("exam_parts").delete().eq("exam_id", exam_id).execute()
    supabase.table("exams").delete().eq("id", exam_id).execute()
    return {"status": "ok"}


class ExamSubmitRequest(BaseModel):
    answers: dict  # {question_id: cau_tra_loi}


@app.post("/exams/{exam_id}/submit")
def submit_exam(exam_id: str, req: ExamSubmitRequest, user=Depends(get_optional_user)):
    require_supabase()
    questions = (
        supabase.table("exam_questions").select("*").eq("exam_id", exam_id)
        .order("question_number").execute().data
    )
    if not questions:
        raise HTTPException(status_code=404, detail="Không tìm thấy đề thi")

    results = []
    correct_count = 0
    section_stats = {}  # {label: {"correct": n, "total": n}}

    for q in questions:
        qid = q["id"]
        user_answer = (req.answers.get(qid) or "").strip()
        correct_answer = (q["correct_answer"] or "").strip()
        is_correct = user_answer.lower() == correct_answer.lower()
        if is_correct:
            correct_count += 1

        label = q.get("section_label") or "Chung"
        section_stats.setdefault(label, {"correct": 0, "total": 0})
        section_stats[label]["total"] += 1
        if is_correct:
            section_stats[label]["correct"] += 1

        results.append({
            "question_id": qid,
            "question_number": q["question_number"],
            "question_text": q["question_text"],
            "your_answer": user_answer,
            "correct_answer": correct_answer,
            "is_correct": is_correct,
            "explanation": q.get("explanation") or "",
            "section_label": q.get("section_label") or "",
        })

    section_breakdown = [
        {"label": label, "correct": s["correct"], "total": s["total"]}
        for label, s in section_stats.items()
    ]

    # Neu da dang nhap, tu luu lai lich su lam bai (khong lam gian doan viec tra ket qua neu loi)
    if user:
        try:
            supabase.table("exam_attempts").insert({
                "exam_id": exam_id,
                "user_id": user["id"],
                "score": correct_count,
                "total": len(questions),
                "answers": req.answers,
                "results": results,
            }).execute()
        except Exception:
            pass

    return {
        "score": correct_count,
        "total": len(questions),
        "results": results,
        "section_breakdown": section_breakdown,
    }


# ---------------------------------------------------------------------------
# Lich su lam bai cua chinh user dang dang nhap
# ---------------------------------------------------------------------------
@app.get("/me/exam-attempts")
def list_my_attempts(user=Depends(get_current_user)):
    attempts = (
        supabase.table("exam_attempts")
        .select("id, exam_id, score, total, created_at, exams(title, skill, is_full_test)")
        .eq("user_id", user["id"])
        .order("created_at", desc=True)
        .execute().data
    )
    return attempts


@app.get("/me/exam-attempts/{attempt_id}")
def get_my_attempt(attempt_id: str, user=Depends(get_current_user)):
    attempt = (
        supabase.table("exam_attempts")
        .select("*, exams(title, skill, is_full_test)")
        .eq("id", attempt_id)
        .single()
        .execute()
    )
    if not attempt.data or attempt.data["user_id"] != user["id"]:
        raise HTTPException(status_code=404, detail="Không tìm thấy lịch sử làm bài")
    return attempt.data


# ---------------------------------------------------------------------------
# Dictation (nghe chép chính tả, tự tách câu từ transcript YouTube - mien phi)
# ---------------------------------------------------------------------------
import re as _re_dictation
from youtube_transcript_api import YouTubeTranscriptApi

WEBSHARE_PROXY_USERNAME = os.environ.get("WEBSHARE_PROXY_USERNAME", "")
WEBSHARE_PROXY_PASSWORD = os.environ.get("WEBSHARE_PROXY_PASSWORD", "")


def _get_youtube_transcript_api():
    """Neu co cau hinh proxy Webshare thi dung, de tranh bi YouTube chan IP cua server cloud (Render)."""
    if WEBSHARE_PROXY_USERNAME and WEBSHARE_PROXY_PASSWORD:
        from youtube_transcript_api.proxies import WebshareProxyConfig
        return YouTubeTranscriptApi(
            proxy_config=WebshareProxyConfig(
                proxy_username=WEBSHARE_PROXY_USERNAME,
                proxy_password=WEBSHARE_PROXY_PASSWORD,
            )
        )
    return YouTubeTranscriptApi()


def extract_youtube_id(url: str) -> str:
    patterns = [
        r"(?:youtu\.be/)([A-Za-z0-9_-]{11})",
        r"(?:v=)([A-Za-z0-9_-]{11})",
        r"(?:embed/)([A-Za-z0-9_-]{11})",
        r"(?:shorts/)([A-Za-z0-9_-]{11})",
    ]
    for p in patterns:
        m = _re_dictation.search(p, url)
        if m:
            return m.group(1)
    # Neu nguoi dung dan thang video ID (11 ky tu, khong co dau /)
    if _re_dictation.fullmatch(r"[A-Za-z0-9_-]{11}", url.strip()):
        return url.strip()
    raise HTTPException(status_code=400, detail="Không nhận diện được video ID từ link YouTube")


_CAPTION_NOISE_RE = _re_dictation.compile(r'\[[^\]]*\]|\([^)]*\)|>>+|♪+', _re_dictation.IGNORECASE)
_MIN_WORDS_PER_SENTENCE = 5


def _clean_caption_text(text: str) -> str:
    """Loai bo cac ky hieu rac cua phu de tu dong: [Music], [Applause], >>, ♪..."""
    text = _CAPTION_NOISE_RE.sub(' ', text)
    text = _re_dictation.sub(r'\s+', ' ', text).strip()
    return text


def _build_word_timeline(transcript):
    timeline = []
    for seg in transcript:
        text = _clean_caption_text(seg["text"].replace("\n", " "))
        if not text:
            continue
        words = text.split()
        if not words:
            continue
        seg_start = seg["start"]
        seg_dur = max(seg["duration"], 0.1)
        per_word = seg_dur / len(words)
        for i, w in enumerate(words):
            word_start = seg_start + i * per_word
            timeline.append({"word": w, "time": word_start, "end": word_start + per_word})
    return timeline


_SENTENCE_END_BUFFER = 0.35  # dem nho, chi du de khong cat cut am cuoi, tranh tran qua cau sau


def _split_into_sentences(timeline, min_words=_MIN_WORDS_PER_SENTENCE):
    sentences = []
    buffer = []
    sent_start = None
    for item in timeline:
        if sent_start is None:
            sent_start = item["time"]
        buffer.append(item)
        if _re_dictation.search(r'[.?!]["\')]*$', item["word"]):
            # Neu cau qua ngan (duoi min_words tu), khong cat - gop tiep vao cau ke tiep
            if len(buffer) < min_words:
                continue
            text = " ".join(w["word"] for w in buffer)
            end_time = item["end"] + _SENTENCE_END_BUFFER  # dung diem KET THUC that cua tu cuoi, khong phai diem bat dau
            sentences.append({"text": text, "start_time": round(sent_start, 2), "end_time": round(end_time, 2)})
            buffer = []
            sent_start = None
    if buffer:
        text = " ".join(w["word"] for w in buffer)
        sentences.append({
            "text": text,
            "start_time": round(sent_start, 2),
            "end_time": round(buffer[-1]["end"] + _SENTENCE_END_BUFFER, 2),
        })
    for i, s in enumerate(sentences):
        s["sentence_number"] = i + 1
    return sentences


class DictationTranscriptRequest(BaseModel):
    youtube_url: str


@app.post("/dictation/fetch-transcript")
def fetch_dictation_transcript(req: DictationTranscriptRequest, user=Depends(get_current_user)):
    require_admin(user)
    video_id = extract_youtube_id(req.youtube_url)
    try:
        ytt_api = _get_youtube_transcript_api()
        fetched = ytt_api.fetch(video_id, languages=["en", "en-US", "en-GB"])
        transcript = fetched.to_raw_data()
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Không lấy được transcript (video có thể không có phụ đề tiếng Anh, hoặc bị chặn phụ đề): {e}",
        )

    if not transcript:
        raise HTTPException(status_code=400, detail="Video này không có transcript")

    timeline = _build_word_timeline(transcript)
    sentences = _split_into_sentences(timeline)
    total_duration = transcript[-1]["start"] + transcript[-1]["duration"]

    return {
        "video_id": video_id,
        "sentences": sentences,
        "duration_seconds": round(total_duration),
        "sentence_count": len(sentences),
    }


class DictationSentenceIn(BaseModel):
    sentence_number: int
    text: str
    start_time: float
    end_time: float


class DictationLessonCreateRequest(BaseModel):
    title: str
    youtube_video_id: str
    level: Optional[str] = None
    sentences: List[DictationSentenceIn]


@app.post("/dictation/lessons")
def create_dictation_lesson(req: DictationLessonCreateRequest, user=Depends(get_current_user)):
    require_admin(user)
    if not req.sentences:
        raise HTTPException(status_code=400, detail="Bài dictation cần có ít nhất 1 câu")

    duration = max((s.end_time for s in req.sentences), default=0)
    lesson_row = {
        "title": req.title,
        "youtube_video_id": req.youtube_video_id,
        "level": req.level,
        "duration_seconds": round(duration),
        "sentence_count": len(req.sentences),
        "created_by": user["id"],
    }
    lesson_res = supabase.table("dictation_lessons").insert(lesson_row).execute()
    lesson_id = lesson_res.data[0]["id"]

    sentence_rows = [{
        "lesson_id": lesson_id,
        "sentence_number": s.sentence_number,
        "text": s.text,
        "start_time": s.start_time,
        "end_time": s.end_time,
    } for s in req.sentences]
    supabase.table("dictation_sentences").insert(sentence_rows).execute()

    return {"status": "ok", "lesson_id": lesson_id}


@app.put("/dictation/lessons/{lesson_id}")
def update_dictation_lesson(lesson_id: str, req: DictationLessonCreateRequest, user=Depends(get_current_user)):
    require_admin(user)
    existing = supabase.table("dictation_lessons").select("id").eq("id", lesson_id).single().execute()
    if not existing.data:
        raise HTTPException(status_code=404, detail="Không tìm thấy bài dictation")

    duration = max((s.end_time for s in req.sentences), default=0)
    supabase.table("dictation_lessons").update({
        "title": req.title,
        "youtube_video_id": req.youtube_video_id,
        "level": req.level,
        "duration_seconds": round(duration),
        "sentence_count": len(req.sentences),
    }).eq("id", lesson_id).execute()

    supabase.table("dictation_sentences").delete().eq("lesson_id", lesson_id).execute()
    sentence_rows = [{
        "lesson_id": lesson_id,
        "sentence_number": s.sentence_number,
        "text": s.text,
        "start_time": s.start_time,
        "end_time": s.end_time,
    } for s in req.sentences]
    supabase.table("dictation_sentences").insert(sentence_rows).execute()

    return {"status": "ok", "lesson_id": lesson_id}


@app.get("/dictation/lessons")
def list_dictation_lessons(level: Optional[str] = None):
    require_supabase()
    q = supabase.table("dictation_lessons").select(
        "id,title,youtube_video_id,level,duration_seconds,sentence_count,created_at"
    ).order("created_at", desc=True)
    if level:
        q = q.eq("level", level)
    return q.execute().data


@app.get("/dictation/lessons/{lesson_id}")
def get_dictation_lesson(lesson_id: str):
    require_supabase()
    lesson = supabase.table("dictation_lessons").select("*").eq("id", lesson_id).single().execute()
    if not lesson.data:
        raise HTTPException(status_code=404, detail="Không tìm thấy bài dictation")
    sentences = (
        supabase.table("dictation_sentences").select("*").eq("lesson_id", lesson_id)
        .order("sentence_number").execute().data
    )
    return {"lesson": lesson.data, "sentences": sentences}


@app.get("/dictation/lessons/{lesson_id}/edit")
def get_dictation_lesson_for_edit(lesson_id: str, user=Depends(get_current_user)):
    require_admin(user)
    lesson = supabase.table("dictation_lessons").select("*").eq("id", lesson_id).single().execute()
    if not lesson.data:
        raise HTTPException(status_code=404, detail="Không tìm thấy bài dictation")
    sentences = (
        supabase.table("dictation_sentences").select("*").eq("lesson_id", lesson_id)
        .order("sentence_number").execute().data
    )
    return {"lesson": lesson.data, "sentences": sentences}


@app.delete("/dictation/lessons/{lesson_id}")
def delete_dictation_lesson(lesson_id: str, user=Depends(get_current_user)):
    require_admin(user)
    supabase.table("dictation_sentences").delete().eq("lesson_id", lesson_id).execute()
    supabase.table("dictation_lessons").delete().eq("id", lesson_id).execute()
    return {"status": "ok"}


class DictationResultIn(BaseModel):
    sentence_number: int
    is_correct: bool
    your_answer: str
    correct_text: str


class DictationSubmitRequest(BaseModel):
    correct_count: int
    total: int
    results: List[DictationResultIn]


@app.post("/dictation/lessons/{lesson_id}/submit")
def submit_dictation(lesson_id: str, req: DictationSubmitRequest, user=Depends(get_optional_user)):
    if user:
        try:
            supabase.table("dictation_attempts").insert({
                "lesson_id": lesson_id,
                "user_id": user["id"],
                "correct_count": req.correct_count,
                "total": req.total,
                "results": [r.dict() for r in req.results],
            }).execute()
        except Exception:
            pass
    return {"status": "ok"}


@app.get("/me/dictation-attempts")
def list_my_dictation_attempts(user=Depends(get_current_user)):
    attempts = (
        supabase.table("dictation_attempts")
        .select("id, lesson_id, correct_count, total, created_at, dictation_lessons(title, level)")
        .eq("user_id", user["id"])
        .order("created_at", desc=True)
        .execute().data
    )
    return attempts
